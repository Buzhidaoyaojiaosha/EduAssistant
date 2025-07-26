import traceback

from flask import Blueprint, render_template, redirect, url_for, flash, request, session, jsonify
from peewee import JOIN
import io

from app.ext import graph
from app.services.course_service import CourseService
from app.services.assignment_service import AssignmentService
from app.services.user_service import UserService
from app.services.knowledge_point_service import KnowledgePointService
from app.services.knowledge_base_service import KnowledgeBaseService
from app.models.user import User
from app.models.course import Course
from app.services.teaching_preparation_service import TeachingPreparationService
from datetime import datetime
from peewee import DoesNotExist
from app.models.knowledge_base import KnowledgeBase
from app.models.course import StudentCourse
from app.models.assignment import Assignment, StudentAssignment
from app.models.NewAdd import Question, StudentAnswer, Feedback, WrongBook, QuestionWrongBook,AIQuestion,AiQuestionStudentAnswer
from app.models.learning_data import (
    KnowledgePoint, AssignmentKnowledgePoint, 
    StudentKnowledgePoint, LearningActivity, 
    KnowledgeBaseKnowledgePoint
)
import re 
from app.utils.logging import logger
import logging
logger = logging.getLogger(__name__)

from flask import current_app, jsonify, request
from app.services.question_generator_service import QuestionGeneratorService

course_bp = Blueprint('course', __name__, url_prefix='/course')

@course_bp.route('/')
def index():
    if 'user_id' not in session:
        return redirect(url_for('auth.login'))
    
    user_id = session['user_id']
    user = User.get_by_id(user_id)
    
    roles = [ur.role.name for ur in user.roles]
    print(f"调试信息 - 用户角色: {roles}")  # 添加调试输出
  
    is_student = 'student' in roles
    is_admin = 'admin' in roles
    is_teacher = 'teacher' in roles
    # 获取搜索关键词
    search_query = request.args.get('q', '').strip()
    
    # 获取课程数据（根据用户角色）
    if UserService.has_role(user, 'teacher'):
        courses = CourseService.get_courses_by_teacher(user_id, search_query)
    else:
        courses = CourseService.get_all_courses(search_query)
    
    return render_template('course/index.html', courses=courses)

@course_bp.route('/create', methods=['GET', 'POST'])
def create():
    if 'user_id' not in session:
        return redirect(url_for('auth.login'))
        
    user_id = session['user_id']
    user = User.get_by_id(user_id)
    
    if not UserService.has_role(user, 'teacher'):
        flash('只有教师可以创建课程。', 'warning')
        return redirect(url_for('course.index'))
    
    if request.method == 'POST':
        name = request.form.get('name')
        code = request.form.get('code')
        description = request.form.get('description')
        
        try:
            course = CourseService.create_course(
                name=name,
                code=code,
                description=description,
                teacher_id=user_id
            )
            flash(f'课程 "{name}" 创建成功!', 'success')
            return redirect(url_for('course.view', course_id=course.id))
        except ValueError as e:
            flash(str(e), 'danger')
    
    return render_template('course/create.html')

@course_bp.route('/<int:course_id>/delete', methods=['GET', 'POST'])
def delete(course_id):
    """删除课程及其所有相关数据"""
    if 'user_id' not in session:
        return redirect(url_for('auth.login'))
    
    user_id = session['user_id']
    
    try:
        course = Course.get_by_id(course_id)
    except Course.DoesNotExist:
        flash('课程不存在。', 'danger')
        return redirect(url_for('course.index'))
    
    # 验证权限 - 只有课程教师可以删除课程
    if course.teacher_id != user_id:
        flash('只有课程教师可以删除课程。', 'warning')
        return redirect(url_for('course.view', course_id=course_id))
    
    if request.method == 'POST':
        try:
            # 使用CourseService删除课程及其所有相关数据
            course_name = course.name
            CourseService.delete_course(course_id)
            flash(f'课程 "{course_name}" 及其所有相关数据已成功删除。', 'success')
            return redirect(url_for('course.index'))
                
        except Exception as e:
            import traceback
            traceback.print_exc()
            flash(f'删除课程失败: {str(e)}', 'danger')
            return redirect(url_for('course.view', course_id=course_id))
    
    # GET请求 - 显示确认删除页面
    return render_template('course/delete_confirm.html', course=course)


# view.html
@course_bp.route('/<int:course_id>')
def view(course_id):
    if 'user_id' not in session:
        return redirect(url_for('auth.login'))
    
    course = Course.get_by_id(course_id)
    user_id = session['user_id']
    user = User.get_by_id(user_id)
    roles = [ur.role.name for ur in user.roles]
    print(f"调试信息 - 用户角色: {roles}")  # 添加调试输出
  
    is_student = 'student' in roles
    is_admin = 'admin' in roles

    # 确认用户是课程的教师或学生
    is_teacher = course.teacher_id == user_id
    is_student = False
    
    if not is_teacher:
        student_courses = CourseService.get_courses_by_student(user_id)
        if course_id in [c.id for c in student_courses]:
            is_student = True
    
    '''if not (is_teacher or is_student):
        flash('您没有访问该课程的权限。', 'warning')
        return redirect(url_for('course.index'))'''
    
    # 获取课程作业
    assignments = AssignmentService.get_course_assignments(course_id)
    
    # 获取课程知识点
    knowledge_points = KnowledgePointService.get_course_knowledge_points(course_id)
    
    # 如果是教师，获取学生列表
    students = None
    if is_teacher or is_admin:
        students = CourseService.get_students_by_course(course_id)
    
    # 如果是学生，获取个人作业情况
    student_assignments = None
    wrong_questions = []
    if is_student:
        student_assignments = AssignmentService.get_student_assignments(user_id, course_id)
        # 新增错题数据查询
        from app.models.NewAdd import StudentAnswer,Question
        from app.models.assignment import Assignment
        
        # 查询当前课程所有作业及关联题目
        course_assignments = (Assignment
                             .select()
                             .where(Assignment.course == course)
                             .order_by(Assignment.due_date.desc()))
        
        for assignment in course_assignments:
            # 查询该作业的错题（答案错误或得分低于题目分值）
            wrong_answers = (StudentAnswer
                            .select()
                            .join(Question)
                            .where(
                                (StudentAnswer.student == user) &
                                (Question.assignment == assignment) &
                                (
                                    (Question.answer != StudentAnswer.commit_answer) |
                                    (StudentAnswer.earned_score < Question.score) |
                                    (StudentAnswer.earned_score.is_null())
                                )
                            ))
            
            if wrong_answers:
                question_data = []
                for answer in wrong_answers:
                    question_data.append({
                        'id': answer.question.question_id,
                        'question_name': answer.question.question_name,
                        'context': answer.question.context,
                        'answer': answer.question.answer,
                        'student_answer': answer.commit_answer,
                        'score': answer.earned_score or 0,
                        'total_points': answer.question.score,
                        'created_at': answer.work_time or datetime.now()
                    })
                
                wrong_questions.append({
                    'assignment': assignment,
                    'questions': question_data
                })
    course = Course.get_by_id(course_id)

    # 获取与该课程相关的知识库条目
    knowledge_entries = CourseService.get_knowledge_base_by_course(course_id)

    return render_template('course/view.html',
                         course=course,
                         is_teacher=is_teacher,
                         is_student=is_student,
                         is_admin=is_admin,
                         assignments=assignments,
                         students=students,
                         student_assignments=student_assignments,
                         knowledge_points=knowledge_points,
                         wrong_questions=wrong_questions,
                         knowledge_entries=knowledge_entries)

@course_bp.route('/<int:course_id>/enroll', methods=['GET', 'POST'])
def enroll(course_id):
    if 'user_id' not in session:
        return redirect(url_for('auth.login'))
    
    user_id = session['user_id']
    user = User.get_by_id(user_id)
    
    if not UserService.has_role(user, 'student'):
        flash('只有学生可以加入课程。', 'warning')
        return redirect(url_for('course.index'))
    
    course = Course.get_by_id(course_id)
    
    if request.method == 'POST':
        try:
            CourseService.enroll_student(course_id, user_id)
            flash(f'成功加入课程 "{course.name}"!', 'success')
        except ValueError as e:
            flash(str(e), 'warning')
        
        return redirect(url_for('course.view', course_id=course_id))
    
    return render_template('course/enroll.html', course=course)

@course_bp.route('/unenroll/<int:course_id>', methods=['POST'])
def unenroll(course_id):
    if 'user_id' not in session:
        return redirect(url_for('auth.login'))
    
    # 检查用户是否已加入该课程
    user_id = session['user_id']
    try:
        if CourseService.unenroll_student(course_id, user_id):
            flash(f'您已成功退出课程', 'success')
        else:
            flash(f'退出课程失败', 'warning')
    except Exception as e:
        flash(str(e), 'warning')
    
    return redirect(url_for('course.view', course_id=course_id))


@course_bp.route('/<int:course_id>/assignment/create', methods=['GET', 'POST'])
def create_assignment(course_id):
    if 'user_id' not in session:
        return redirect(url_for('auth.login'))
    
    user_id = session['user_id']
    course = Course.get_by_id(course_id)
    
    # 验证权限
    if course.teacher_id != user_id:
        flash('只有课程教师可以创建作业。', 'warning')
        return redirect(url_for('course.view', course_id=course_id))
    
    if request.method == 'POST':
        title = request.form.get('title')
        description = request.form.get('description')
        due_date = datetime.fromisoformat(request.form.get('due_date'))
        total_points = float(request.form.get('total_points', 100))
        
        assignment = AssignmentService.create_assignment(
            title=title,
            description=description,
            course_id=course_id,
            due_date=due_date,
            total_points=total_points
        )
        
        # 自动分配给所有学生
        assigned_count = AssignmentService.assign_to_students(assignment.id)
        
        flash(f'作业已创建并分配给{assigned_count}名学生。', 'success')
        return redirect(url_for('course.view_assignment', assignment_id=assignment.id))
    
    return render_template('course/create_assignment.html', course=course)

@course_bp.route('/assignment/<int:assignment_id>/delete', methods=['POST'])
def delete_assignment(assignment_id):
    """删除作业及其所有相关数据"""
    if 'user_id' not in session:
        flash('请先登录', 'warning')
        return redirect(url_for('auth.login'))
    
    try:
        assignment = Assignment.get_by_id(assignment_id)
        course_id = assignment.course.id
        user_id = session['user_id']
        
        # 验证权限 - 只有课程教师可以删除作业
        if assignment.course.teacher_id != user_id:
            flash('只有课程教师可以删除作业', 'warning')
            return redirect(url_for('course.view_assignment', assignment_id=assignment_id))
        
        # 使用事务确保数据一致性
        with Assignment._meta.database.atomic():
            # 1. 获取作业相关的所有题目
            questions = Question.select().where(Question.assignment == assignment)
            
            # 2. 删除所有题目相关的数据
            if questions:
                question_ids = [q.question_id for q in questions]
                
                # 删除学生答案
                StudentAnswer.delete().where(
                    StudentAnswer.question_id.in_(question_ids)
                ).execute()
                
                # 删除错题本关联
                QuestionWrongBook.delete().where(
                    QuestionWrongBook.question_id.in_(question_ids)
                ).execute()

                # 获取AI生成的题目ID列表
                ai_questions = AIQuestion.select().where(
                    AIQuestion.original_question.in_(question_ids)
                )
                
                if ai_questions:
                    ai_question_ids = [ai_q.ai_question_id for ai_q in ai_questions]
                    
                    # 删除AI生成题目的学生答案
                    AiQuestionStudentAnswer.delete().where(
                        AiQuestionStudentAnswer.ai_question.in_(ai_question_ids)
                    ).execute()

                # 删除AI生成的题目
                AIQuestion.delete().where(
                    AIQuestion.original_question.in_(question_ids)
                ).execute()
                
                # 删除题目本身
                Question.delete().where(
                    Question.assignment == assignment
                ).execute()
            
            # 3. 删除学生作业记录
            StudentAssignment.delete().where(
                StudentAssignment.assignment == assignment
            ).execute()
            
            # 4. 删除作业反馈
            Feedback.delete().where(
                Feedback.assignment == assignment
            ).execute()
            
            # 5. 删除作业与知识点的关联
            AssignmentKnowledgePoint.delete().where(
                AssignmentKnowledgePoint.assignment == assignment
            ).execute()
            
            # 6. 最后删除作业本身
            assignment_title = assignment.title
            assignment.delete_instance()
            
            flash(f'作业 "{assignment_title}" 及其所有相关数据已成功删除', 'success')
        
        return redirect(url_for('course.view', course_id=course_id))
    
    except DoesNotExist:
        flash('作业不存在', 'danger')
        return redirect(url_for('dashboard.index'))
    except Exception as e:
        import traceback
        traceback.print_exc()
        flash(f'删除作业失败: {str(e)}', 'danger')
        return redirect(url_for('course.view_assignment', assignment_id=assignment_id))
    
#查看作业详情页
@course_bp.route('/assignment/<int:assignment_id>')
def view_assignment(assignment_id):
    if 'user_id' not in session:
        return redirect(url_for('auth.login'))
    
    from app.models.assignment import Assignment, StudentAssignment
    from app.models.NewAdd import Question, Feedback
    
    assignment = Assignment.get_by_id(assignment_id)
    user_id = session['user_id']
    
    # 检查权限
    is_teacher = assignment.course.teacher_id == user_id
    is_admin = UserService.has_role(User.get_by_id(user_id), 'admin')
    print(f"调试信息 - 用户角色: {is_teacher}, {is_admin}")  # 添加调试输出
    student_assignment = None
    feedback = None  # 新增feedback变量
    
    # 获取作业关联的所有题目
    questions = Question.select(
        Question.question_id,
        Question.question_name,
        Question.context,
        Question.answer,
        Question.score,
        Question.status).where(Question.assignment == assignment).order_by(Question.status)
        
    # 如果是学生，获取学生作业状态和评语
    if not is_teacher:
        student_assignment = StudentAssignment.get_or_none(
            StudentAssignment.student_id == user_id,
            StudentAssignment.assignment_id == assignment_id
        )
        
        # 获取该学生的评语
        if student_assignment:
            feedback = Feedback.get_or_none(
                (Feedback.assignment == assignment) &
                (Feedback.student_id == user_id))
    
    # 如果是教师，获取所有学生提交情况
    submissions = None
    if is_teacher:
        submissions = (StudentAssignment
                      .select()
                      .where(StudentAssignment.assignment_id == assignment_id))
    
    # 获取作业关联的知识点
    knowledge_points = KnowledgePointService.get_assignment_knowledge_points(assignment_id)
    
    # 添加当前时间变量
    now = datetime.now()
    
    # 获取AI生成的相似题目（仅教师可见）
    ai_questions = []
    if is_teacher:
        ai_questions = AIQuestion.select().where(
            (AIQuestion.assignment == assignment) 
        ).order_by(AIQuestion.created_time.desc())
    
    return render_template('course/view_assignment.html',
                         is_admin = is_admin,
                         assignment=assignment,
                         questions=questions,
                         ai_questions=ai_questions,
                         is_teacher=is_teacher,
                         student_assignment=student_assignment,
                         submissions=submissions,
                         knowledge_points=knowledge_points,
                         now=now,
                         feedback=feedback)  


@course_bp.route('/ai_question/<int:question_id>/edit', methods=['GET', 'POST'])
def edit_ai_question(question_id):
    ai_question = AIQuestion.get_or_none(question_id)
    
    if request.method == 'POST':
        try:
            # 获取表单数据并更新
            ai_question.question_name = request.form.get('question_name', ai_question.question_name)
            ai_question.answer = request.form.get('answer', ai_question.answer)
            ai_question.analysis = request.form.get('analysis', ai_question.analysis)
            ai_question.status = int(request.form.get('status', ai_question.status))
            
            # 处理题目内容
            if ai_question.status == 1:  # 选择题
                options = request.form.getlist('options[]')
                question_stem = request.form.get('question_stem', '')
                
                # 为每个选项添加字母前缀 (A., B., C., ...)
                prefixed_options = []
                for i, option in enumerate(options):
                    letter = chr(65 + i)  # A, B, C, ...
                    prefixed_options.append(f"{letter}. {option.strip()}")
                
                # 合并题干和选项
                ai_question.context = question_stem + '\n' + '\n'.join(prefixed_options)
                
            elif ai_question.status == 2:  # 判断题
                ai_question.context = request.form.get('context', '')
            else:  # 简答题或编程题
                ai_question.context = request.form.get('context', '')
            
            ai_question.save()
            flash('AI题目修改成功', 'success')
            return redirect(url_for('course.view_assignment', assignment_id=ai_question.assignment.id))
        except Exception as e:
            flash(f'修改失败: {str(e)}', 'danger')
    
    # 准备编辑表单数据
    question_stem = ""
    options = []
    if ai_question.status == 1:  # 如果是选择题
        parts = ai_question.context.split('\n')
        question_stem = parts[0] if parts else ""
        options = []
        
        # 提取选项内容（去掉字母前缀）
        for part in parts[1:]:
            # 匹配 "A. 选项内容" 格式
            if re.match(r'^[A-Z]\.\s', part):
                options.append(part[3:].strip())  # 去掉前3个字符（如"A. "）
            else:
                options.append(part.strip())
    
    return render_template('course/edit_ai_question.html',
                         ai_question=ai_question,
                         question_stem=question_stem,
                         options=options,
                         question_types={
                             1: '选择题',
                             2: '判断题', 
                             3: '简答题',
                             4: '编程题'
                         })

# 添加审核路由
@course_bp.route('/ai_question/<int:question_id>/approve', methods=['POST'])
def approve_ai_question(question_id):
    try:
        question = AIQuestion.get_by_id(question_id)
        question.is_approved = True
        question.save()
        flash('AI题目已通过审核', 'success')
    except Exception as e:
        flash(f'审核失败: {str(e)}', 'danger')
    return redirect(url_for('course.view_assignment', assignment_id=question.assignment.id))

# 添加移除审核路由 若老师审查不通过，直接删除
@course_bp.route('/ai_question/<int:question_id>/unapprove', methods=['POST'])
def unapprove_ai_question(question_id):
    try:
        question = AIQuestion.get_by_id(question_id)
        assignment_id = question.assignment.id
        question.delete_instance()  # 直接删除记录
        flash('此AI题目已删除', 'warning')
    except Exception as e:
        flash(f'删除失败: {str(e)}', 'danger')
    return redirect(url_for('course.view_assignment', assignment_id=assignment_id))


#学生提交作业
@course_bp.route('/assignment/<int:assignment_id>/submit', methods=['POST'])
def submit_assignment(assignment_id):
    if 'user_id' not in session:
        return redirect(url_for('auth.login'))
    
    user_id = session['user_id']
    
    try:
        from app.models.NewAdd import StudentAnswer, Question, Feedback  # 添加Feedback导入
        from app.models.assignment import Assignment, StudentAssignment
        from app.models.course import StudentCourse
        from app.services.assignment_service import AssignmentService  # 导入AssignmentService
        
        # 获取作业
        assignment = Assignment.get_by_id(assignment_id)
        
        # 验证学生是否有权限提交此作业
        enrollment = StudentCourse.get_or_none(
            (StudentCourse.student_id == user_id) &
            (StudentCourse.course_id == assignment.course_id)
        )
        
        if not enrollment:
            flash('您没有权限提交此作业。', 'danger')
            return redirect(url_for('course.view_assignment', assignment_id=assignment_id))
        
        # 查找或创建学生作业记录
        student_assignment, created = StudentAssignment.get_or_create(
            student_id=user_id,
            assignment_id=assignment_id
        )
        
        # 检查是否已经提交过
        if student_assignment.status >= 1:
            flash('作业已经提交过，无法重复提交。', 'warning')
            return redirect(url_for('course.view_assignment', assignment_id=assignment_id))
        
        # 获取作业的所有题目
        questions = Question.select().where(Question.assignment_id == assignment_id)
        
        if not questions.exists():
            flash('此作业暂无题目，无法提交。', 'warning')
            return redirect(url_for('course.view_assignment', assignment_id=assignment_id))
        
        # 处理每个题目的答案
        total_score = 0
        answered_count = 0
        answer_details = []  # 用于收集答案详情供AI生成评语
        
        for question in questions:
            answer_key = f'answer_{question.question_id}'
            
            if answer_key in request.form:
                student_answer = request.form.get(answer_key, '').strip()
                
                if not student_answer:
                    continue  # 跳过空答案
                
                answered_count += 1
                # 计算得分（选择题和判断题可以自动评分）
                earned_score = 0
                if question.status == 1:  # 选择题
                    correct_answer = question.answer.strip()
                    if student_answer == correct_answer:
                        earned_score = question.score
                elif question.status == 2:  # 判断题
                    # 转换存储格式：1->"对", 0->"错"
                    stored_answer = "对" if student_answer == "1" or student_answer == "对" else "错"
                    correct_answer = question.answer.strip()
                   
                    print(f"Correct answer for question {question.question_id}: {correct_answer}\n")
                    print(f"Student answer for question {question.question_id}: {stored_answer}\n")
                    if stored_answer == correct_answer:
                        earned_score = question.score
                    student_answer = stored_answer  # 使用转换后的答案进行存储
                else:  # 简答题和编程题
                    try:
                        earned_score = AssignmentService.grade_short_answer_with_deepseek(
                            question=question.context,
                            student_answer=student_answer,
                            max_score=question.score
                        )
                    except Exception as e:
                        print(f"AI评分失败: {str(e)}")
                        earned_score = 0
                
                # 存储学生回答
                StudentAnswer.create(
                    student_id=user_id,
                    question_id=question.question_id,
                    commit_answer=student_answer,
                    earned_score=earned_score
                )
                
                total_score += earned_score
                answer_details.append({
                    'question': question.context,  
                    'score': earned_score,
                    'max_score': question.score,
                    'answer': student_answer  
                })
        
        if answered_count == 0:
            flash('请至少回答一道题目。', 'warning')
            return redirect(url_for('course.view_assignment', assignment_id=assignment_id))
        
        # 更新学生作业状态
        student_assignment.work_time = datetime.now()
        student_assignment.status = 1  # 已提交
        student_assignment.final_score = total_score if total_score > 0 else None
        student_assignment.save()
        
        # 调用AI生成评语
        try:
            feedback_content = AssignmentService.generate_feedback_with_deepseek(
                student_answers=answer_details,
                total_score=total_score,
                max_score=assignment.total_points
            )
            
            # 存储评语到Feedback表
            Feedback.create(
                assignment=assignment,
                student_id=user_id,
                comment=feedback_content
            )
        except Exception as e:
            print(f"生成评语失败: {str(e)}")
            # 即使评语生成失败也不影响作业提交
        
        flash(f'作业已成功提交！回答了 {answered_count} 道题目，得分: {total_score}/{assignment.total_points}', 'success')
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        flash(f'提交作业失败: {str(e)}', 'danger')
    
    return redirect(url_for('course.view_assignment', assignment_id=assignment_id))

#查看提交的作业的作业详情
@course_bp.route('/assignment/<int:assignment_id>/submission/<int:student_id>', methods=['GET'])
def view_submission(assignment_id, student_id):
    if 'user_id' not in session:
        return redirect(url_for('auth.login'))
    
    user_id = session['user_id']
    
    assignment = AssignmentService.get_assignment_by_id(assignment_id)
    from app.models.user import User
    from app.models.assignment import StudentAssignment
    from app.models.NewAdd import StudentAnswer, Question
    
    # 检查权限
    is_teacher = assignment.course.teacher_id == user_id
    is_student = user_id == student_id
    
    if not (is_teacher or is_student):
        flash('您没有权限查看此提交。', 'warning')
        return redirect(url_for('dashboard.index'))
    
    student = User.get_by_id(student_id)
    submission = StudentAssignment.get_or_none(
        StudentAssignment.student==student,
        StudentAssignment.assignment==assignment
    )
    
    if not submission:
        flash('未找到该提交记录。', 'warning')
        return redirect(url_for('course.view_assignment', assignment_id=assignment_id))
    
    # 获取学生的所有答案，按题目顺序排序
    student_answers = (StudentAnswer
                      .select(StudentAnswer, Question)
                      .join(Question)
                      .where(
                          (StudentAnswer.student_id == student_id) &
                          (Question.assignment_id == assignment_id)
                      )
                      .order_by(Question.question_id))

    # 获取评语
    from app.models.NewAdd import Feedback
    feedback = Feedback.get_or_none(
        (Feedback.assignment == assignment) &
        (Feedback.student_id == student_id))
    
    return render_template('course/view_submission.html', 
                          assignment=assignment, 
                          student=student,
                          submission=submission,
                          student_answers=student_answers,
                          is_teacher=is_teacher,
                          feedback=feedback)  # 添加feedback参数

@course_bp.route('/assignment/<int:assignment_id>/grade/<int:student_id>', methods=['POST'])
def grade_assignment(assignment_id, student_id):
    if 'user_id' not in session:
        return redirect(url_for('auth.login'))
    
    user_id = session['user_id']
    from app.models.assignment import Assignment
    
    assignment = Assignment.get_by_id(assignment_id)
    
    # 验证权限
    if assignment.course.teacher_id != user_id:
        flash('只有教师可以评分。', 'warning')
        return redirect(url_for('dashboard.index'))
    
    score = float(request.form.get('score', 0))
    feedback = request.form.get("feedback")
    
    try:
        # 更新为使用最终得分
        from app.models.assignment import StudentAssignment
        student_assignment = StudentAssignment.get(
            StudentAssignment.student_id == student_id,
            StudentAssignment.assignment_id == assignment_id
        )
        student_assignment.final_score = score  # 使用final_score而不是score
        student_assignment.feedback = feedback
        student_assignment.status = 2  # 已批改
        student_assignment.save()
        
        flash('评分已保存。', 'success')
    except Exception as e:
        flash(f'评分失败: {str(e)}', 'danger')
    
    return redirect(url_for('course.view_submission', assignment_id=assignment_id, student_id=student_id))

@course_bp.route('/<int:course_id>/knowledge_point/add', methods=['POST'])
def add_knowledge_point(course_id):
    """添加新的知识点到课程中"""
    if 'user_id' not in session:
        return redirect(url_for('auth.login'))
    
    user_id = session['user_id']
    course = Course.get_by_id(course_id)
    
    # 验证权限
    if course.teacher_id != user_id:
        flash('只有课程教师可以添加知识点', 'warning')
        return redirect(url_for('course.view', course_id=course_id))
    
    name = request.form.get('name')
    description = request.form.get('description', '')
    parent_id = request.form.get('parent_id')
    
    # 如果父级ID为空字符串，则设为None
    if parent_id == '':
        parent_id = None
    elif parent_id:
        parent_id = int(parent_id)
    
    try:
        KnowledgePointService.add_knowledge_to_graph(
            name=name,
            course_id=course_id,
            description=description,
            parent_id=parent_id
        )
        flash(f'知识点 "{name}" 创建成功!', 'success')
    except ValueError as e:
        flash(str(e), 'danger')
    
    return redirect(url_for('course.view', course_id=course_id))

@course_bp.route('/<int:course_id>/knowledge_point/edit', methods=['POST'])
def edit_knowledge_point(course_id):
    """编辑课程中的知识点"""
    if 'user_id' not in session:
        return redirect(url_for('auth.login'))
    
    user_id = session['user_id']
    course = Course.get_by_id(course_id)
    
    # 验证权限
    if course.teacher_id != user_id:
        flash('只有课程教师可以编辑知识点', 'warning')
        return redirect(url_for('course.view', course_id=course_id))
    
    knowledge_point_id = int(request.form.get('knowledge_point_id'))
    name = request.form.get('name')
    description = request.form.get('description', '')
    parent_id = request.form.get('parent_id')
    
    # 如果父级ID为空字符串，则设为None
    if parent_id == '':
        parent_id = None
    elif parent_id:
        parent_id = int(parent_id)
    
    try:
        # 获取知识点实例
        knowledge_point = KnowledgePointService.get_knowledge_point(knowledge_point_id)
        
        # 检查知识点是否属于当前课程
        if knowledge_point.course_id != course_id:
            raise ValueError('该知识点不属于当前课程')
        
        # 防止循环引用
        if parent_id and parent_id == knowledge_point_id:
            raise ValueError('知识点不能以自己作为父级')
        
        # 更新知识点
        knowledge_point.name = name
        knowledge_point.description = description
        knowledge_point.parent_id = parent_id
        knowledge_point.save()

        try:
            KnowledgePointService.update_knowledge_point_node(
                kp_id=knowledge_point.id,
                name=name,
                description=description,
                parent_id=parent_id,
                course_id=course_id
            )
        except Exception as e:
            flash(f'图数据库同步失败：{e}', 'warning')
        
        flash(f'知识点 "{name}" 更新成功!', 'success')
    except ValueError as e:
        flash(str(e), 'danger')
    
    return redirect(url_for('course.view', course_id=course_id))

@course_bp.route('/<int:course_id>/knowledge_point/delete', methods=['POST'])
def delete_knowledge_point(course_id):
    """删除课程中的知识点"""
    if 'user_id' not in session:
        return redirect(url_for('auth.login'))
    
    user_id = session['user_id']
    course = Course.get_by_id(course_id)
    
    # 验证权限
    if course.teacher_id != user_id:
        flash('只有课程教师可以删除知识点', 'warning')
        return redirect(url_for('course.view', course_id=course_id))
    
    knowledge_point_id = int(request.form.get('knowledge_point_id'))
    
    try:
        # 获取知识点实例
        from app.models.learning_data import KnowledgePoint
        knowledge_point = KnowledgePointService.get_knowledge_point(knowledge_point_id)
        
        # 检查知识点是否属于当前课程
        if knowledge_point.course_id != course_id:
            raise ValueError('该知识点不属于当前课程')
        
        # 检查是否有子知识点
        children = KnowledgePoint.select().where(KnowledgePoint.parent_id == knowledge_point_id)
        if children.count() > 0:
            raise ValueError('该知识点有子知识点，请先删除或重新分配子知识点')
        
        # 保存名称以供通知
        kp_name = knowledge_point.name
        
        # 删除知识点
        knowledge_point.delete_instance()
        flash(f'知识点 "{kp_name}" 已删除', 'success')

        try:
            KnowledgePointService.delete_knowledge_point_node_from_graph(kp_name)
        except Exception as e:
            flash(f'图数据库同步失败（删除知识点）: {e}', 'warning')
        

    except ValueError as e:
        flash(str(e), 'danger')
    
    return redirect(url_for('course.view', course_id=course_id))


#从Excel文件中导入知识点
@course_bp.route('/<int:course_id>/import_knowledge_points', methods=['POST'])
def import_knowledge_points(course_id):
    if 'user_id' not in session:
        return redirect(url_for('auth.login'))
    
    user = User.get_by_id(session['user_id'])
    course = Course.get_by_id(course_id)
    
    # 验证用户是否为该课程的教师
    if course.teacher_id != user.id:
        flash('只有课程教师可以导入知识点。', 'warning')
        return redirect(url_for('course.view', course_id=course_id))
    
    if 'excel_file' not in request.files:
        flash('未选择文件。', 'warning')
        return redirect(url_for('course.view', course_id=course_id))
    
    file = request.files['excel_file']
    if file.filename == '':
        flash('未选择文件。', 'warning')
        return redirect(url_for('course.view', course_id=course_id))
    
    if file:
        try:
            # 读取文件内容到内存
            file_content = file.read()
            file_stream = io.BytesIO(file_content)
        
        # 先导入到PostgreSQL，获取名称到ID的映射
            id_cache = KnowledgePointService.import_excel_to_knowledge_points(file_stream, course_id)
        
        # 重置文件流指针
            file_stream.seek(0)
        
        # 再导入到Neo4j，传递ID映射
            KnowledgePointService.excel_to_knowledge_point_graph(file_stream, course_id, id_cache)
           
            flash('知识点导入成功。', 'success')
        except Exception as e:
            flash(f'知识点导入失败: {str(e)}', 'danger')
            traceback.print_exc()
        
        return redirect(url_for('course.view', course_id=course_id))



#查看课程知识图谱
@course_bp.route('/<int:course_id>/view_knowledge_graph')
def get_course_graph(course_id):
    query = """
    MATCH (k1:Knowledge {course_id: $course_id})
    OPTIONAL MATCH (k1)-[r]->(k2:Knowledge {course_id: $course_id})
    RETURN k1, r, k2
    """
    results = graph.execute_query(query, course_id=course_id, database="neo4j")

    nodes = {}
    edges = []

    for records in results:
        for record in records :

            k1, r, k2 = record.values()
            for node in [k1, k2]:
                if node is not None:
                    nid = getattr(node, "id", None) or getattr(node, "element_id", None)
                    if nid not in nodes:
                        nodes[nid] = {
                            "id": nid,
                            "label": node.get("name", "未知"),
                            "group": node.get("level", 1)
                        }

            if r is not None and k1 is not None and k2 is not None:
                edges.append({
                    "from": getattr(k1, "id", None) or getattr(k1, "element_id", None),
                    "to": getattr(k2, "id", None) or getattr(k2, "element_id", None),
                    "label": getattr(r, "type", "关联")
                })

        return jsonify({
            "nodes": list(nodes.values()),
            "edges": edges
        })


@course_bp.route('/assignment/<int:assignment_id>/knowledge_points', methods=['GET', 'POST'])
def assignment_knowledge_points(assignment_id):
    """管理作业关联的知识点"""
    if 'user_id' not in session:
        return redirect(url_for('auth.login'))
    
    from app.models.assignment import Assignment
    
    assignment = Assignment.get_by_id(assignment_id)
    course_id = assignment.course_id
    user_id = session['user_id']
    
    # 验证权限
    if assignment.course.teacher_id != user_id:
        flash('只有课程教师可以管理作业知识点', 'warning')
        return redirect(url_for('course.view_assignment', assignment_id=assignment_id))
    
    if request.method == 'POST':
        # 获取提交的知识点ID列表
        knowledge_point_ids = request.form.getlist('knowledge_point_ids', type=int)
        
        # 获取权重
        weights = {}
        for kp_id in knowledge_point_ids:
            weight = request.form.get(f'weight_{kp_id}', 1.0, type=float)
            weights[kp_id] = weight
        
        try:
            # 清除旧的关联并添加新的
            from app.models.learning_data import AssignmentKnowledgePoint
            AssignmentKnowledgePoint.delete().where(
                AssignmentKnowledgePoint.assignment_id == assignment_id
            ).execute()
            
            if knowledge_point_ids:
                KnowledgePointService.add_knowledge_points_to_assignment(
                    assignment_id, knowledge_point_ids, weights
                )
            
            flash('作业知识点关联已更新', 'success')
        except ValueError as e:
            flash(str(e), 'danger')
        
        return redirect(url_for('course.view_assignment', assignment_id=assignment_id))
    
    # 获取课程所有知识点
    course_knowledge_points = KnowledgePointService.get_course_knowledge_points(course_id)
    course_knowledge_points = [kp for kp in course_knowledge_points if kp is not None]  # 过滤掉None
    # 获取作业已关联的知识点
    assignment_knowledge_points = KnowledgePointService.get_assignment_knowledge_points(assignment_id)
    
    return render_template('course/assignment_knowledge_points.html',
                          assignment=assignment,
                          course_knowledge_points=course_knowledge_points,
                          assignment_knowledge_points=assignment_knowledge_points)

@course_bp.route('/assignment/<int:assignment_id>/add-question', methods=['GET', 'POST'])
def add_question(assignment_id):
    from app.models.assignment import Assignment
    from app.models.course import Course
    from app.services.assignment_service import AssignmentService  # 导入AssignmentService
    import logging
    logger = logging.getLogger(__name__)
    
    assignment = Assignment.get_by_id(assignment_id)
    course = assignment.course
    
    # 验证权限
    user_id = session['user_id']
    if assignment.course.teacher_id != user_id:
        flash('只有课程教师可以添加题目', 'warning')
        return redirect(url_for('course.view_assignment', assignment_id=assignment_id))
    
    if request.method == 'POST':
        try:
            # 获取表单数据
            question_type = int(request.form.get('type', 0))
            answer = request.form.get('answer', '').strip()
            
            # 如果是判断题，转换答案格式
            if question_type == 2:  # 判断题
                if answer == '1':
                    answer = '对'
                elif answer == '0':
                    answer = '错'  # 修正了之前的错误，0应该对应"错"
                else:
                    # 如果输入的不是1或0，但题目类型是判断题，给出错误提示
                    flash('判断题答案只能是1(对)或0(错)', 'danger')
                    return render_template('course/add_question.html', assignment=assignment)
            
            print(f"Processed answer for question type {question_type}: {answer}\n")
            
            # 创建题目
            question = Question.create(
                assignment=assignment,
                course=course,
                question_name=request.form.get('name'),
                context=request.form.get('context'),
                answer=answer,  # 使用转换后的答案
                analysis=request.form.get('analysis'),
                score=float(request.form.get('score', 10.0)),
                status=question_type
            )
            
            # 更新作业总分
            AssignmentService.update_assignment_total_points(assignment_id)
            
            # # 添加成功后才生成相似题目
            # try:
            #     # 调用生成相似题目的函数
            #     generated_questions = AssignmentService.generate_similar_questions_with_ai(
            #         original_question=question,
            #         assignment=assignment,
            #         num_questions=3
            #     )
                
            #     if generated_questions:
            #         print(f'题目添加成功，并已生成{len(generated_questions)}道相似题目，请到AI题目管理中审核', 'success')
            #     else:
            #         print('题目添加成功，但生成相似题目失败', 'warning')
            # except Exception as ai_error:
            #     logger.error(f"生成相似题目失败: {str(ai_error)}")
            #     print('题目添加成功，但生成相似题目时出错', 'warning')
            print(f'题目添加成功，ID: {question.question_id}', 'success')
            return redirect(url_for('course.view_assignment', assignment_id=assignment_id))
        
        except Exception as e:
            flash(f'添加题目失败: {str(e)}', 'danger')
    
    return render_template('course/add_question.html', assignment=assignment)

@course_bp.route('/question/<int:question_id>/edit', methods=['GET', 'POST'])
def edit_question(question_id):
    question = Question.get_by_id(question_id)
    user_id = session['user_id']
    
    # 验证权限
    if question.assignment.course.teacher_id != user_id:
        flash('只有课程教师可以更新题目', 'warning')
        return redirect(url_for('course.view_assignment', assignment_id=question.assignment_id))
    
    if request.method == 'POST':
        try:
            # 保存旧的分数用于比较
            old_score = question.score
            
            # 获取表单数据并更新
            question.question_name = request.form.get('name', question.question_name)
            question.answer = request.form.get('answer', question.answer)
            question.analysis = request.form.get('analysis', question.analysis)
            question.status = int(request.form.get('type', question.status))
            question.score = float(request.form.get('score', question.score))
            
            # 处理题目内容
            if question.status == 1:  # 选择题
                options = request.form.getlist('options[]')
                question_stem = request.form.get('context', '')
                
                # 为每个选项添加字母前缀 (A., B., C., ...)
                prefixed_options = []
                for i, option in enumerate(options):
                    letter = chr(65 + i)  # A, B, C, ...
                    prefixed_options.append(f"{letter}. {option.strip()}")
                
                # 合并题干和选项
                question.context = question_stem + '\n' + '\n'.join(prefixed_options)
                
            elif question.status == 2:  # 判断题
                question.context = request.form.get('context', '')
            else:  # 简答题或编程题
                question.context = request.form.get('context', '')
            
            question.save()
            
            # 如果分数有变化，更新作业总分
            if old_score != question.score:
                AssignmentService.update_assignment_total_points(question.assignment.id)
            
            flash('题目更新成功', 'success')
            return redirect(url_for('course.view_assignment', assignment_id=question.assignment.id))
            
        except Exception as e:
            flash(f'更新题目失败: {str(e)}', 'danger')
    
    # 准备编辑表单数据
    question_stem = ""
    options = []
    if question.status == 1:  # 如果是选择题
        parts = question.context.split('\n')
        question_stem = parts[0] if parts else ""
        options = []
        
        # 提取选项内容（去掉字母前缀）
        for part in parts[1:]:
            # 匹配 "A. 选项内容" 格式
            if re.match(r'^[A-Z]\.\s', part):
                options.append(part[3:].strip())  # 去掉前3个字符（如"A. "）
            else:
                options.append(part.strip())
    
    return render_template('course/edit_question.html',
                         question=question,
                         question_stem=question_stem,
                         options=options,
                         question_types={
                             1: '选择题',
                             2: '判断题',
                             3: '简答题',
                             4: '编程题'
                         })

@course_bp.route('/question/<int:question_id>/delete', methods=['POST'])
def delete_question(question_id):
    """删除题目并更新作业总分（同时删除关联的AI题目）"""
    if 'user_id' not in session:
        return redirect(url_for('auth.login'))
    
    user_id = session['user_id']
    import logging
    logger = logging.getLogger(__name__)    
    try:
        question = Question.get_by_id(question_id)
        assignment_id = question.assignment.id
        
        # 验证权限
        if question.assignment.course.teacher_id != user_id:
            flash('只有课程教师可以删除题目', 'warning')
            return redirect(url_for('course.view_assignment', assignment_id=assignment_id))
        
        # 删除关联的AI题目（新增部分）
        ai_questions = AIQuestion.select().where(
            (AIQuestion.original_question == question_id) 
           
        )
        for ai_question in ai_questions:
            ai_question.delete_instance()
        
        # 删除原题目
        question.delete_instance()
        
        # 更新作业总分
        AssignmentService.update_assignment_total_points(assignment_id)
        
        flash('题目及关联AI题目已成功删除', 'success')  # 修改提示信息
    except DoesNotExist:
        flash('题目不存在', 'danger')
    except Exception as e:
        flash(f'删除题目失败: {str(e)}', 'danger')
        logger.error(f"删除题目{question_id}失败: {str(e)}", exc_info=True)
    
    return redirect(url_for('course.view_assignment', assignment_id=assignment_id))

@course_bp.route('/assignment/<int:assignment_id>/grade_student/<int:student_id>', methods=['GET'])
def grade_student_assignment(assignment_id, student_id):
    if 'user_id' not in session:
        return redirect(url_for('auth.login'))
    
    user_id = session['user_id']
    
    from app.models.assignment import Assignment, StudentAssignment
    from app.models.NewAdd import StudentAnswer, Question
    from app.models.user import User
    
    assignment = Assignment.get_by_id(assignment_id)
    
    # 验证权限
    if assignment.course.teacher_id != user_id:
        flash('只有教师可以评分作业。', 'warning')
        return redirect(url_for('dashboard.index'))
    
    student = User.get_by_id(student_id)
    
    # 获取学生作业记录
    student_assignment = StudentAssignment.get_or_none(
        StudentAssignment.student_id == student_id,
        StudentAssignment.assignment_id == assignment_id
    )
    
    if not student_assignment or student_assignment.status == 0:
        flash('该学生尚未提交作业。', 'warning')
        return redirect(url_for('course.view_assignment', assignment_id=assignment_id))
    
    # 获取学生的所有答案
    student_answers = StudentAnswer.select().join(Question).where(
        (StudentAnswer.student_id == student_id) &
        (Question.assignment_id == assignment_id)
    )
    
    # 获取评语
    feedback = Feedback.get_or_none(
        (Feedback.assignment_id == assignment_id) &
        (Feedback.student_id == student_id)
    )
    
    return render_template('course/grade_assignment.html', 
                          assignment=assignment,
                          student=student,
                          student_assignment=student_assignment,
                          student_answers=student_answers,
                          feedback=feedback)  # 传递feedback对象到模板

# 老师再次评分和写评语
@course_bp.route('/assignment/<int:assignment_id>/grade_student/<int:student_id>', methods=['POST'])
def grade_student_answers(assignment_id, student_id):
    if 'user_id' not in session:
        return redirect(url_for('auth.login'))
    
    user_id = session['user_id']
    
    from app.models.assignment import Assignment, StudentAssignment
    from app.models.NewAdd import StudentAnswer, Question, Feedback  # 添加Feedback导入
    
    assignment = Assignment.get_by_id(assignment_id)
    
    # 验证权限
    if assignment.course.teacher_id != user_id:
        flash('只有教师可以评分作业。', 'warning')
        return redirect(url_for('dashboard.index'))
    
    try:
        # 1. 处理每道题的评分
        student_answers = (StudentAnswer
                          .select(StudentAnswer, Question)
                          .join(Question)
                          .where(
                              (StudentAnswer.student_id == student_id) &
                              (Question.assignment_id == assignment_id)
                          ))
        
        total_score = 0
        for answer in student_answers:
            score_key = f'score_{answer.submission_id}'
            if score_key in request.form:
                score = float(request.form.get(score_key, 0))
                answer.earned_score = score
                answer.save()
                total_score += score
        
        # 2. 更新学生作业记录
        student_assignment = StudentAssignment.get(
            StudentAssignment.student_id == student_id,
            StudentAssignment.assignment_id == assignment_id
        )
        student_assignment.final_score = total_score
        student_assignment.status = 2  # 已批改
        student_assignment.save()
        
        # 3. 处理评语（使用Feedback类）
        feedback_text = request.form.get('feedback', '').strip()
        if feedback_text:  # 只有评语不为空时才处理
            # 查找是否已有评语记录
            try:
                feedback = Feedback.get(
                    (Feedback.assignment == assignment) &
                    (Feedback.student_id == student_id)
                )
                # 更新现有评语
                feedback.comment = feedback_text
                feedback.save()
            except Feedback.DoesNotExist:
                # 创建新评语记录
                Feedback.create(
                    assignment=assignment,
                    student_id=student_id,
                    comment=feedback_text
                )
        
        flash('评分和评语已保存。', 'success')
    except Exception as e:
        flash(f'评分失败: {str(e)}', 'danger')
    
    return redirect(url_for('course.view_assignment', assignment_id=assignment_id))
@course_bp.route('/<int:course_id>/knowledge_entries', methods=['GET'])
def get_knowledge_entries(course_id):
    """获取课程知识库条目（用于前端选择）"""
    try:
        knowledge_entries = CourseService.get_knowledge_base_by_course(course_id)
        
        # 格式化返回数据
        formatted_entries = []
        for entry in knowledge_entries:
            # 创建内容预览（截断长文本）
            content_preview = entry.content
            if not entry.content.startswith(('http://', 'https://', 'ftp://')):
                content_preview = entry.content[:150] + '...' if len(entry.content) > 150 else entry.content
            else:
                content_preview = "链接资源"
            
            formatted_entries.append({
                'id': entry.id,
                'title': entry.title,
                'category': entry.category or '未分类',
                'type': '可下载资源' if entry.content.startswith(('http://', 'https://', 'ftp://')) else '纯文本资源',
                'content_preview': content_preview,
                'created_at': entry.created_at.strftime('%Y-%m-%d')
            })
        
        return jsonify(formatted_entries)
    
    except Exception as e:
        logger.error(f"获取知识库条目失败: {str(e)}")
        return jsonify({'error': '获取知识库内容失败'}), 500
    
@course_bp.route('/api/generate_teaching_outline', methods=['POST'])
def api_generate_teaching_outline():
    try:
        # 检查用户是否登录
        if 'user_id' not in session:
            return jsonify({'error': '用户未登录，请先登录'}), 401
        
        # 检查请求内容类型
        if not request.is_json:
            return jsonify({'error': '请求格式错误，需要JSON格式'}), 400
        
        # 获取请求数据
        data = request.get_json()
        course_id = data['course_id']
        material_type = data.get('material_type', 'syllabus')  # 新增：大纲或PPT
        selected_knowledge_ids = data.get('selected_knowledge_ids', [])

        if not data:
            return jsonify({'error': '请求数据为空'}), 400
            
        if 'course_id' not in data:
            return jsonify({'error': '缺少必要的参数: course_id'}), 400
        
        # 验证course_id格式
        try:
            course_id = int(course_id)
        except (ValueError, TypeError):
            return jsonify({'error': '课程ID格式错误'}), 400
        
        # 验证课程存在
        course = Course.get_by_id(course_id)
        if not course:
            return jsonify({'error': '课程不存在'}), 404
        
        # 调用教学准备服务生成大纲（传递选中的知识库ID）
        try:
            result = TeachingPreparationService.generate_outline(
                course_id, 
                material_type=material_type,  # 传递材料类型
                selected_knowledge_ids=selected_knowledge_ids
            )
            
        except Exception as service_error:
            logger.error(f"TeachingPreparationService.generate_outline 失败: {str(service_error)}")
            return jsonify({'error': '教学大纲生成服务暂时不可用，请稍后重试'}), 500
        
        # 检查生成结果
        if not result:
            return jsonify({'error': '生成服务返回空结果'}), 500
            
        if 'error' in result:
            logger.error(f"生成教学大纲失败 - 课程ID: {course_id}, 错误: {result['error']}")
            return jsonify({'error': result['error']}), 500
        
        # 验证返回数据的完整性
        required_fields = ['content', 'file_base64', 'filename', 'title', 'download_ready']
        for field in required_fields:
            if field not in result:
                logger.error(f"生成结果缺少必要字段: {field}")
                return jsonify({'error': f'生成结果不完整，缺少: {field}'}), 500
        
        # 返回成功结果
        return jsonify({
            'success': True,
            'content': result['content'],
            'file_base64': result['file_base64'],  # 统一字段名
            'filename': result['filename'],
            'title': result['title'],
            'download_ready': result['download_ready'],
            'material_type': result.get('material_type', 'syllabus'),  # 新增
            'file_type': result.get('file_type', 'pdf')  # 新增
        }), 200
        
    except Exception as e:
        logger.error(f"API生成教学大纲异常: {str(e)}", exc_info=True)
        return jsonify({'error': '服务器内部错误，请稍后重试'}), 500
    
@course_bp.route('/<int:course_id>/generate_assessment', methods=['POST'])
def generate_assessment(course_id):
    """生成考核题目接口"""
    try:
        data = request.get_json()
        question_settings = data.get('question_settings', {})  # 获取题目设置
        selected_knowledge_ids = data.get('selected_knowledge_ids', [])
        
        # 调用生成函数，传递题目设置
        questions = QuestionGeneratorService.generate_questions_with_ai(
            course_id, 
            question_settings,
            selected_knowledge_ids=selected_knowledge_ids
        )
        
        return jsonify({
            'success': True,
            'questions': questions
        })
    except Exception as e:
        current_app.logger.error(f"生成考核题目失败: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@course_bp.route('/<int:course_id>/assignments', methods=['GET'])
def get_assignments(course_id):
    """获取课程的作业列表"""
    try:
        # 验证课程是否存在
        try:
            course = Course.get_by_id(course_id)
        except DoesNotExist:
            return jsonify({
                'success': False,
                'error': '课程不存在'
            }), 404
        
        # 获取作业列表
        assignments = Assignment.select().where(Assignment.course == course)
        
        # 构造返回数据
        assignments_data = []
        for assignment in assignments:
            assignments_data.append({
                'id': assignment.id,
                'title': assignment.title,
                'description': assignment.description,
                'due_date': assignment.due_date.strftime('%Y-%m-%d') if assignment.due_date else None,
                'total_points': float(assignment.total_points) if assignment.total_points else 0.0
            })
        
        return jsonify({
            'success': True,
            'assignments': assignments_data
        })
        
    except Exception as e:
        current_app.logger.error(f"获取作业列表失败: {str(e)}")
        return jsonify({
            'success': False,
            'error': '获取作业列表失败'
        }), 500
    
@course_bp.route('/<int:course_id>/save_assessment', methods=['POST'])
def save_assessment(course_id):
    """保存生成的考核题目"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({
                'success': False,
                'error': '请求数据为空'
            }), 400
            
        assignment_id = data.get('assignment_id')
        questions = data.get('questions', [])
        
        if not questions:
            return jsonify({
                'success': False,
                'error': '没有题目数据'
            }), 400
        
        # 验证课程是否存在
        try:
            course = Course.get_by_id(course_id)
        except DoesNotExist:
            return jsonify({
                'success': False,
                'error': '课程不存在'
            }), 404
        
        # 处理保存逻辑
        if assignment_id == "new":
            # 创建新作业
            assignment = Assignment.create(
                title=f"AI生成考核题目-{datetime.now().strftime('%Y%m%d')}",
                description=f"通过AI自动生成的考核题目，生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                course=course,  # 传递Course对象，而不是course_id
                due_date=datetime.now().replace(hour=23, minute=59, second=59),  # 设置为当天23:59:59
                total_points=0.0
            )
            assignment_id = assignment.id
        else:
            # 验证作业是否存在
            try:
                assignment = Assignment.get_by_id(assignment_id)
            except DoesNotExist:
                return jsonify({
                    'success': False,
                    'error': '作业不存在'
                }), 404
        
        # 保存题目
        total_points = assignment.total_points
        created_questions = []
        
        for question_data in questions:
            # 验证题目数据
            if not question_data.get('content') or not question_data.get('answer'):
                continue
            
            # 根据题目类型设置分数和状态
            question_type = question_data.get('type', '简答题')
            if question_type == '简答题':
                score = 10.0
                status = 3  # 简答题
            elif question_type == '编程题':
                score = 10.0
                status = 4  # 编程题
            elif question_type == '判断题':
                score = 5.0
                status = 2  # 判断题
            elif question_type == '选择题':
                score = 5.0
                status = 1  # 选择题
            else:
                score = 5.0
                status = 3  # 默认为简答题
            
            # 创建题目对象
            new_question = Question.create(
                question_name=question_data.get('title', f"题目{len(created_questions) + 1}"),
                assignment=assignment,
                course=course,
                context=question_data['content'],
                answer=question_data['answer'],
                analysis=question_data.get('analysis', ''),
                score=score,
                status=status
            )
            
            total_points += score
            created_questions.append(new_question)
        
        # 更新作业总分
        assignment.total_points = total_points
        assignment.save()
        
        return jsonify({
            'success': True,
            'assignment_id': assignment_id,
            'questions_created': len(created_questions),
            'total_points': total_points
        })
        
    except Exception as e:
        current_app.logger.error(f"保存考核题目失败: {str(e)}")
        return jsonify({
            'success': False,
            'error': '保存失败，请稍后重试'
        }), 500

@course_bp.route('/assignment/<int:assignment_id>/edit_due_date', methods=['GET', 'POST'])
def edit_assignment(assignment_id):
    """编辑作业"""
    if 'user_id' not in session:
        flash('请先登录', 'warning')
        return redirect(url_for('auth.login'))
    
    try:
        assignment = Assignment.get_by_id(assignment_id)
        user_id = session['user_id']
        
        # 验证权限 - 只有课程教师可以修改
        if assignment.course.teacher_id != user_id:
            flash('只有课程教师可以编辑作业', 'warning')
            return redirect(url_for('course.view_assignment', assignment_id=assignment_id))
        
        if request.method == 'POST':
            # 获取表单数据
            new_title = request.form.get('title')
            new_due_date_str = request.form.get('new_due_date')
            # 将字符串转换为datetime对象
            new_due_date = datetime.fromisoformat(new_due_date_str)

            # 验证数据
            if not new_title or not new_due_date_str:
                flash('作业名称和截止时间不能为空', 'danger')
                return redirect(url_for('course.edit_assignment', assignment_id=assignment_id))
            
            # 更新作业的截止时间
            AssignmentService.update_title(assignment_id,new_title)
            # 更新作业的截止时间
            AssignmentService.update_due_date(assignment_id,new_due_date)
            
            flash('作业信息已更新', 'success')
            return redirect(url_for('course.view_assignment', assignment_id=assignment_id))
        
        # GET请求 - 显示修改表单
        # 将datetime对象转换为字符串格式，用于填充表单
        current_due_date = assignment.due_date.strftime('%Y-%m-%dT%H:%M')
        return render_template('course/edit_due_date.html', 
                             assignment=assignment,
                             current_due_date=current_due_date)
    
    except DoesNotExist:
        flash('作业不存在', 'danger')
        return redirect(url_for('dashboard.index'))
    except ValueError as e:
        flash(f'日期格式错误: {str(e)}', 'danger')
        return redirect(url_for('course.edit_assignment_due_date', assignment_id=assignment_id))
    except Exception as e:
        flash(f'更新失败: {str(e)}', 'danger')
        return redirect(url_for('course.view_assignment', assignment_id=assignment_id))
    
@course_bp.route('/question/<int:question_id>/generate_similar', methods=['POST'])
def generate_similar_questions(question_id):
    """为指定题目生成相似题目"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'error': '用户未登录'}), 401
    
    try:
        # 获取题目和用户信息
        question = Question.get_by_id(question_id)
        user_id = session['user_id']
        
        # 验证权限
        if question.assignment.course.teacher_id != user_id:
            return jsonify({
                'success': False,
                'error': '只有课程教师可以生成相似题目'
            }), 403
        
        # 调用生成函数
        generated_questions = AssignmentService.generate_similar_questions_with_ai(
            original_question=question,
            assignment=question.assignment,
            num_questions=3
        )
        
        if generated_questions:
            flash(f'成功生成 {len(generated_questions)}道相似题目', 'success')
        else:
            flash(f'生成相似题目失败，请检查题目内容', 'danger')

        return redirect(url_for('course.view_assignment', assignment_id=question.assignment.id))
            
    except DoesNotExist:
        flash(f'生成相似题目失败，题目不存在', 'danger')
        return redirect(url_for('course.view_assignment', assignment_id=question.assignment.id))
        
    except Exception as e:
        flash(f'生成相似题目失败，请检查题目内容', 'danger')
        return redirect(url_for('course.view_assignment', assignment_id=question.assignment.id))

# 添加保存教学资料的路由
@course_bp.route('/save_teaching_material', methods=['POST'])
def save_teaching_material():
    try:
        data = request.json
        result = TeachingPreparationService.save_teaching_material(
            course_id=data['course_id'],
            title=data['title'],
            file_base64=data['file_base64'],
            filename=data['filename']
        )
        
        if result['success']:
            return jsonify({
                "success": True,
                "message": "教学资料已保存至知识库"
            })
        else:
            return jsonify({
                "success": False,
                "error": result.get('error', '保存失败')
            }), 400
            
    except Exception as e:
        current_app.logger.error(f"保存教学资料失败: {str(e)}")
        return jsonify({
            "success": False,
            "error": "服务器内部错误"
        }), 500
    


# 导出作业
from io import BytesIO
from docx import Document
from flask import send_file
from datetime import datetime

@course_bp.route('/<int:course_id>/assignment/<int:assignment_id>/export')
def export_assignment(assignment_id, course_id):
    if 'user_id' not in session:
        return redirect(url_for('auth.login'))
    
    assignment = Assignment.get_by_id(assignment_id)
    questions = Question.select().where(Question.assignment == assignment).order_by(Question.status)
    
    user = User.get_by_id(session['user_id'])
    print(f"调试信息 - 正在删除用户: {user.username}, ID: {user.id}")
    roles = [ur.role.name for ur in user.roles]
    print(f"调试信息 - 用户角色: {roles}")  # 添加调试输出

    is_admin = 'admin' in roles
    
    print(f"调试信息 - 用户角色: {roles}")  # 添加调试输出

    try:
        # 创建Word文档
        doc = Document()
        
        # 添加作业标题和基本信息
        doc.add_heading(assignment.title, level=1)
        doc.add_paragraph(f'课程: {assignment.course.name} ({assignment.course.code})')
        doc.add_paragraph(f'截止日期: {assignment.due_date.strftime("%Y-%m-%d %H:%M")}')
        doc.add_paragraph(f'总分: {assignment.total_points}')
        if assignment.description:
            doc.add_paragraph(f'描述: {assignment.description}')
        
        # 添加题目
        doc.add_heading('题目列表', level=2)
        for idx, question in enumerate(questions, start=1):
            doc.add_heading(f'题目{idx}: {question.question_name} ({question.score}分)', level=3)
            
            # 处理题目内容
            if question.status == 1:  # 选择题
                lines = question.context.split('\n')
                doc.add_paragraph(lines[0])  # 题目描述
                for option in lines[1:]:
                    if option.strip():
                        doc.add_paragraph(option, style='List Bullet')
            else:  # 其他类型题目
                doc.add_paragraph(question.context)
            
            # 如果是教师，显示参考答案
            if assignment.course.teacher_id == session['user_id'] or is_admin:
                doc.add_paragraph(f'参考答案: {question.answer}', style='Intense Quote')
        
        # 保存到内存
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # 返回文件
        filename = f"{assignment.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    except Exception as e:
        flash(f'导出作业失败: {str(e)}', 'danger')
        return redirect(url_for('course.view_assignment', assignment_id=assignment_id))